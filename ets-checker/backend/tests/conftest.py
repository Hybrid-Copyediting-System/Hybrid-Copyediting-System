from __future__ import annotations

from pathlib import Path

import pytest

FIXTURES_DIR = Path(__file__).parent / "fixtures"


@pytest.fixture
def fixtures_dir() -> Path:
    return FIXTURES_DIR


def pytest_configure(config: pytest.Config) -> None:
    """Generate synthetic fixture .docx files that exercise specific rules.

    Fixtures are created once and reused; if they already exist they are left
    untouched so hand-crafted files can override the generated ones.
    """
    FIXTURES_DIR.mkdir(exist_ok=True)
    _make_broken_margins()
    _make_broken_abstract()


def _make_broken_margins() -> None:
    path = FIXTURES_DIR / "broken_margins.docx"
    if path.exists():
        return
    from docx import Document
    from docx.shared import Cm

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(4.0)
        section.bottom_margin = Cm(4.0)
        section.left_margin = Cm(4.0)
        section.right_margin = Cm(4.0)
    doc.add_paragraph("Abstract", style="Heading 1")
    doc.add_paragraph("This is the abstract text for testing margin validation.")
    doc.add_paragraph("Keywords: test, fixture, margins")
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Body text goes here for margin checking purposes.")
    doc.add_paragraph("References", style="Heading 1")
    doc.add_paragraph("Smith, J. (2020). A test paper. Journal, 1(1), 1–10.")
    doc.save(str(path))


def _make_broken_abstract() -> None:
    path = FIXTURES_DIR / "broken_abstract.docx"
    if path.exists():
        return
    from docx import Document

    doc = Document()
    doc.add_paragraph("Abstract", style="Heading 1")
    # 260 distinct words — exceeds the 250-word limit
    filler = " ".join(f"word{i}" for i in range(260))
    doc.add_paragraph(filler)
    doc.add_paragraph("Keywords: test, fixture, abstract")
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Body text goes here.")
    doc.add_paragraph("References", style="Heading 1")
    doc.add_paragraph("Smith, J. (2020). A test paper. Journal, 1(1), 1–10.")
    doc.save(str(path))
