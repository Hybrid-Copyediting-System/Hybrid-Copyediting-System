from __future__ import annotations

import io
import json
from pathlib import Path

import pytest
from docx import Document as _Document
from httpx import ASGITransport, AsyncClient

from ets_checker.server import app


def _build_synthetic_docx() -> io.BytesIO:
    """Smallest valid .docx that exercises every rule path without skips."""
    buf = io.BytesIO()
    doc = _Document()
    doc.add_paragraph("Title")
    doc.add_paragraph("Abstract", style="Heading 1")
    doc.add_paragraph("This is the abstract text.")
    doc.add_paragraph("Keywords: test")
    doc.add_paragraph("Introduction", style="Heading 1")
    doc.add_paragraph("Body text.")
    doc.add_paragraph("References", style="Heading 1")
    doc.add_paragraph("Smith, J. (2020). Test paper. Journal, 1(1), 1.")
    doc.save(buf)
    buf.seek(0)
    return buf


def _parse_sse_events(body: str) -> list[tuple[str, dict]]:
    """Decode a server-sent-events stream body into a list of (event, data) pairs."""
    events: list[tuple[str, dict]] = []
    for chunk in body.split("\n\n"):
        if not chunk.strip():
            continue
        event_name = ""
        data_str = ""
        for line in chunk.splitlines():
            if line.startswith("event: "):
                event_name = line[len("event: "):].strip()
            elif line.startswith("data: "):
                data_str = line[len("data: "):]
        if event_name:
            events.append((event_name, json.loads(data_str) if data_str else {}))
    return events


def _get_fixture(name: str) -> Path:
    p = Path(__file__).parent / "fixtures" / name
    if not p.exists():
        pytest.skip(f"Fixture {name} not found")
    return p


@pytest.mark.asyncio
async def test_health() -> None:
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        r = await ac.get("/api/health")
    assert r.status_code == 200
    assert r.json() == {"status": "ok"}


@pytest.mark.asyncio
async def test_check_no_file() -> None:
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        r = await ac.post("/api/check")
    assert r.status_code == 422 or r.status_code == 400


@pytest.mark.asyncio
async def test_check_wrong_extension() -> None:
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        r = await ac.post(
            "/api/check",
            files={"file": ("test.txt", io.BytesIO(b"hello"), "text/plain")},
        )
    assert r.status_code == 400
    # The error must explain *why* — otherwise a future change that returns
    # a 400 for a different reason (e.g. CSRF) would silently pass this test.
    detail = r.json().get("detail", "").lower()
    assert "docx" in detail or "extension" in detail


@pytest.mark.asyncio
async def test_check_doc_extension() -> None:
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        r = await ac.post(
            "/api/check",
            files={"file": ("test.doc", io.BytesIO(b"hello"), "application/msword")},
        )
    assert r.status_code == 400
    assert "docx" in r.json()["detail"].lower()


@pytest.mark.asyncio
async def test_check_annotated_wrong_extension() -> None:
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        r = await ac.post(
            "/api/check/annotated",
            files={"file": ("test.txt", io.BytesIO(b"hello"), "text/plain")},
        )
    assert r.status_code == 400
    detail = r.json().get("detail", "").lower()
    assert "docx" in detail or "extension" in detail


@pytest.mark.asyncio
async def test_check_annotated_returns_docx() -> None:
    """Round-trip: build a synthetic docx, POST it, expect a docx back."""
    buf = _build_synthetic_docx()

    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        r = await ac.post(
            "/api/check/annotated",
            files={
                "file": (
                    "synthetic.docx",
                    buf,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    assert r.status_code == 200
    assert r.headers["content-type"].startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert "synthetic.annotated.docx" in r.headers["content-disposition"]
    assert r.content[:2] == b"PK"


@pytest.mark.asyncio
async def test_check_template() -> None:
    p = _get_fixture("ets_template.docx")
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        with open(p, "rb") as f:
            r = await ac.post(
                "/api/check",
                files={"file": ("ets_template.docx", f, "application/vnd.openxmlformats-officedocument.wordprocessingml.document")},
            )
    from ets_checker.rules.runner import _REGISTRY

    assert r.status_code == 200
    data = r.json()
    assert "summary" in data
    expected_total = len(_REGISTRY.sync_rules) + len(_REGISTRY.async_rules)
    assert data["summary"]["total_checks"] == expected_total


# ── /api/check/stream ────────────────────────────────────────────────


@pytest.mark.asyncio
async def test_stream_emits_progress_and_complete() -> None:
    """The SSE stream must emit progress events and exactly one final
    ``complete`` event whose payload deserialises into a CheckReport."""
    from ets_checker.models import CheckReport
    from ets_checker.rules.runner import _REGISTRY

    buf = _build_synthetic_docx()
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test", timeout=30.0) as ac:
        r = await ac.post(
            "/api/check/stream",
            files={
                "file": (
                    "synthetic.docx", buf,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    assert r.status_code == 200
    assert r.headers["content-type"].startswith("text/event-stream")

    events = _parse_sse_events(r.text)
    names = [name for name, _ in events]
    # Must start with a "parsing" progress and end with exactly one "complete".
    assert names[0] == "progress"
    assert events[0][1].get("phase") == "parsing"
    assert names.count("complete") == 1
    assert names[-1] == "complete"
    # No "error" events on the happy path.
    assert "error" not in names

    # Each registered sync rule must produce a "rule" progress event.
    rule_events = [
        data for name, data in events
        if name == "progress" and data.get("phase") == "rule"
    ]
    rule_ids_seen = {e["rule_id"] for e in rule_events}
    registered_sync = {rid for rid, *_ in _REGISTRY.sync_rules}
    assert registered_sync.issubset(rule_ids_seen)

    # Final report deserialises and includes every registered rule.
    final = events[-1][1]
    report = CheckReport.model_validate(final)
    expected_total = len(_REGISTRY.sync_rules) + len(_REGISTRY.async_rules)
    assert report.summary.total_checks == expected_total


@pytest.mark.asyncio
async def test_stream_invalid_docx_emits_error_event() -> None:
    """A garbled .docx must surface as an ``error`` SSE event, not a 500."""
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as ac:
        r = await ac.post(
            "/api/check/stream",
            files={
                "file": (
                    "bad.docx", io.BytesIO(b"not a real docx"),
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
        )
    assert r.status_code == 200
    events = _parse_sse_events(r.text)
    names = [n for n, _ in events]
    assert "error" in names
    assert "complete" not in names


# ── /api/check/annotated with report_json ────────────────────────────


@pytest.mark.asyncio
async def test_annotated_uses_provided_report_json() -> None:
    """When a report is provided, the endpoint must reuse it instead of
    re-running checks. We verify by sending a single-finding report and
    counting comments in the returned docx."""
    import zipfile
    from datetime import datetime, timezone
    from io import BytesIO

    from lxml import etree

    from ets_checker.models import (
        CheckDetail,
        CheckReport,
        CheckResult,
        Locator,
        ReportSummary,
    )

    report = CheckReport(
        file_name="synthetic.docx",
        timestamp=datetime.now(timezone.utc),
        summary=ReportSummary(total_checks=1, passed=0, errors=1, warnings=0, info=0),
        results=[
            CheckResult(
                rule_id="custom.marker",
                category="Test",
                name="Marker",
                status="fail",
                severity="error",
                details=[
                    CheckDetail(
                        location="paragraph 0",
                        locator=Locator(kind="paragraph", paragraph_index=0),
                        message="UNIQUE_REPORT_JSON_MARKER",
                    ),
                ],
            ),
        ],
    )

    buf = _build_synthetic_docx()
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test", timeout=30.0) as ac:
        r = await ac.post(
            "/api/check/annotated",
            files={
                "file": (
                    "synthetic.docx", buf,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"report_json": report.model_dump_json()},
        )
    assert r.status_code == 200

    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    with zipfile.ZipFile(BytesIO(r.content)) as zf:
        comments_xml = zf.read("word/comments.xml").decode("utf-8")
    com_root = etree.fromstring(comments_xml.encode("utf-8"))
    comments = com_root.findall(f"{{{W_NS}}}comment")

    # Exactly one comment, and the unique marker must be inside it — proving
    # the supplied report (not a fresh run) drove the annotation.
    assert len(comments) == 1
    assert "UNIQUE_REPORT_JSON_MARKER" in comments_xml


@pytest.mark.asyncio
async def test_annotated_invalid_report_json_falls_back() -> None:
    """Garbage in ``report_json`` should not 500; the endpoint should ignore
    it and fall back to running checks fresh."""
    buf = _build_synthetic_docx()
    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test", timeout=30.0) as ac:
        r = await ac.post(
            "/api/check/annotated",
            files={
                "file": (
                    "synthetic.docx", buf,
                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                )
            },
            data={"report_json": "{not valid json"},
        )
    assert r.status_code == 200
    assert r.content[:2] == b"PK"
