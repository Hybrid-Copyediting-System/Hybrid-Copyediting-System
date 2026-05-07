"""Edge-case fixtures the production parser/runner must handle gracefully:

- A genuinely empty docx (no body paragraphs).
- A corrupted upload (not a real ZIP).
- A document with no detectable sections.
- A document where every paragraph is inside a table.

These tests ensure the rule pipeline never crashes on degenerate input
and that the appropriate error surfaces (DocumentParseError for invalid
input, regular CheckResults for content-shaped degeneracies).
"""
from __future__ import annotations

import asyncio
import io
from pathlib import Path

import pytest
from docx import Document

from ets_checker.parser.docx_parser import parse
from ets_checker.rules.runner import _REGISTRY, run_async
from ets_checker.services import DocumentParseError, parse_document


def _save(doc: Document, path: Path) -> Path:
    doc.save(str(path))
    return path


# ─── empty docx ─────────────────────────────────────────────────────


class TestEmptyDocument:
    def test_empty_docx_parses(self, tmp_path: Path) -> None:
        """python-docx always emits at least one empty paragraph; the parser
        must still produce a ParsedDocument with empty section/citation lists."""
        path = _save(Document(), tmp_path / "empty.docx")
        doc = parse(str(path))
        assert doc.sections == []
        assert doc.citations == []
        assert doc.references == []

    def test_empty_docx_runs_all_rules(self, tmp_path: Path) -> None:
        path = _save(Document(), tmp_path / "empty.docx")
        doc = parse(str(path))
        report = asyncio.run(run_async(doc, "empty.docx"))
        expected = len(_REGISTRY.sync_rules) + len(_REGISTRY.async_rules)
        assert report.summary.total_checks == expected
        # Required-sections rule must report all three missing.
        req = next(
            r for r in report.results if r.rule_id == "structure.required_sections"
        )
        assert req.status == "fail"
        assert len(req.details) == 3


# ─── corrupted docx ─────────────────────────────────────────────────


class TestCorruptedDocument:
    def test_invalid_zip_raises_parse_error(self, tmp_path: Path) -> None:
        path = tmp_path / "corrupt.docx"
        path.write_bytes(b"this is not a real docx file at all")
        with pytest.raises(DocumentParseError):
            parse_document(str(path))

    def test_truncated_zip_raises_parse_error(self, tmp_path: Path) -> None:
        # A real docx truncated to a few bytes — header looks ZIP-ish but the
        # rest is missing. Should still surface as a parse error, not a crash.
        buf = io.BytesIO()
        Document().save(buf)
        path = tmp_path / "trunc.docx"
        path.write_bytes(buf.getvalue()[:50])
        with pytest.raises(DocumentParseError):
            parse_document(str(path))


# ─── no detected sections ──────────────────────────────────────────


class TestNoSections:
    def test_plain_paragraphs_no_sections(self, tmp_path: Path) -> None:
        """A document with paragraphs but no headings should produce zero
        sections — and the rule pipeline must not crash on that."""
        doc = Document()
        for i in range(5):
            doc.add_paragraph(f"Plain paragraph {i}.")
        path = _save(doc, tmp_path / "no_sections.docx")

        parsed = parse(str(path))
        assert parsed.sections == []

        report = asyncio.run(run_async(parsed, "no_sections.docx"))
        # All rules must run; required-sections fails.
        expected = len(_REGISTRY.sync_rules) + len(_REGISTRY.async_rules)
        assert report.summary.total_checks == expected


# ─── all paragraphs inside tables ──────────────────────────────────


class TestAllInTable:
    def test_only_table_content_does_not_crash(self, tmp_path: Path) -> None:
        doc = Document()
        table = doc.add_table(rows=3, cols=2)
        for r in range(3):
            for c in range(2):
                table.cell(r, c).text = f"cell {r},{c}"
        path = _save(doc, tmp_path / "all_in_table.docx")

        parsed = parse(str(path))
        # Every parsed paragraph that has content should be marked is_in_table.
        non_empty = [p for p in parsed.paragraphs if p.text.strip()]
        if non_empty:
            assert all(p.is_in_table for p in non_empty)

        report = asyncio.run(run_async(parsed, "all_in_table.docx"))
        expected = len(_REGISTRY.sync_rules) + len(_REGISTRY.async_rules)
        assert report.summary.total_checks == expected


# ─── file-size boundary ────────────────────────────────────────────


class TestFileSizeBoundary:
    """Verify the 50 MB ``MAX_FILE_SIZE`` boundary in services.save_temp.

    Generating a real 50 MB file per test run is wasteful, so we monkey-patch
    ``MAX_FILE_SIZE`` to a small value and assert the boundary behaviour.
    """

    @pytest.mark.asyncio
    async def test_payload_at_limit_does_not_413(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        from httpx import ASGITransport, AsyncClient

        from ets_checker import services
        from ets_checker.server import app

        limit = 1024
        monkeypatch.setattr(services, "MAX_FILE_SIZE", limit)

        # Junk payload exactly at the limit — save_temp uses ``> MAX_FILE_SIZE``,
        # so this must be accepted by the size check (then rejected later as
        # an unparseable docx, surfacing as 422 — anything but 413 is fine).
        payload = b"X" * limit
        transport = ASGITransport(app=app)
        async with AsyncClient(transport=transport, base_url="http://test") as ac:
            r = await ac.post(
                "/api/check",
                files={
                    "file": (
                        "atlimit.docx", io.BytesIO(payload),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
        assert r.status_code != 413

    @pytest.mark.asyncio
    async def test_payload_over_limit_returns_413(
        self, monkeypatch: pytest.MonkeyPatch
    ) -> None:
        from httpx import ASGITransport, AsyncClient

        from ets_checker import services
        from ets_checker.server import app

        limit = 1024
        monkeypatch.setattr(services, "MAX_FILE_SIZE", limit)

        payload = b"X" * (limit + 1)
        transport = ASGITransport(app=app)
        async with AsyncClient(transport=transport, base_url="http://test") as ac:
            r = await ac.post(
                "/api/check",
                files={
                    "file": (
                        "toobig.docx", io.BytesIO(payload),
                        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    )
                },
            )
        assert r.status_code == 413
        assert "too large" in r.json()["detail"].lower()
