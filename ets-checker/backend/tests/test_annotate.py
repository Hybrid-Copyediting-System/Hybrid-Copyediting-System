from __future__ import annotations

import zipfile
from datetime import datetime, timezone
from io import BytesIO
from pathlib import Path

import pytest
from docx import Document
from lxml import etree

from ets_checker.exporter import annotate
from ets_checker.exporter.anchor import build_paragraph_element_index
from ets_checker.models import (
    CheckDetail,
    CheckReport,
    CheckResult,
    Locator,
    ReportSummary,
)
from ets_checker.parser.docx_parser import parse

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
NSMAP = {"w": W_NS}


def _ids_at(root: etree._Element, xpath: str) -> set[str]:
    out: set[str] = set()
    for el in root.iterfind(xpath):
        cid = el.get(f"{{{W_NS}}}id")
        if cid is not None:
            out.add(cid)
    return out


@pytest.fixture
def synthetic_docx(tmp_path: Path) -> Path:
    """Build a minimal valid .docx with body paragraphs and one table."""
    doc = Document()
    doc.add_paragraph("Title of the paper")
    doc.add_paragraph("Abstract")
    doc.add_paragraph("This is the abstract text.")
    doc.add_paragraph("Introduction body text goes here.")
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "header A"
    table.cell(0, 1).text = "header B"
    table.cell(1, 0).text = "row A"
    table.cell(1, 1).text = "row B"
    doc.add_paragraph("Conclusion paragraph.")
    out = tmp_path / "synthetic.docx"
    doc.save(str(out))
    return out


def _two_finding_report() -> CheckReport:
    return CheckReport(
        file_name="synthetic.docx",
        timestamp=datetime.now(timezone.utc),
        summary=ReportSummary(
            total_checks=2,
            passed=0,
            errors=1,
            warnings=1,
            info=0,
        ),
        results=[
            CheckResult(
                rule_id="layout.margins",
                category="Layout",
                name="Margin check",
                status="fail",
                severity="error",
                details=[
                    CheckDetail(
                        location="document",
                        locator=Locator(kind="document"),
                        message="Top margin does not match ET&S",
                        expected=2.54,
                        actual=3.0,
                    ),
                ],
            ),
            CheckResult(
                rule_id="structure.abstract_length",
                category="Structure",
                name="Abstract length check",
                status="fail",
                severity="warning",
                details=[
                    CheckDetail(
                        location="Abstract",
                        locator=Locator(kind="paragraph", paragraph_index=2),
                        message="Abstract too long",
                        expected="≤ 200 words",
                        actual="500 words",
                    ),
                ],
            ),
        ],
    )


class TestAnchorIndexParity:
    def test_index_length_matches_parser(self, synthetic_docx: Path) -> None:
        parsed = parse(str(synthetic_docx))
        document = Document(str(synthetic_docx))
        index = build_paragraph_element_index(document)
        assert len(index) == len(parsed.paragraphs)


class TestAnnotate:
    def test_returns_valid_docx_bytes(self, synthetic_docx: Path) -> None:
        blob = annotate(str(synthetic_docx), _two_finding_report())
        assert blob[:2] == b"PK"
        Document(BytesIO(blob))

    def test_source_file_unchanged(self, synthetic_docx: Path) -> None:
        before = synthetic_docx.read_bytes()
        annotate(str(synthetic_docx), _two_finding_report())
        after = synthetic_docx.read_bytes()
        assert before == after

    def test_contains_comments_part(self, synthetic_docx: Path) -> None:
        blob = annotate(str(synthetic_docx), _two_finding_report())
        with zipfile.ZipFile(BytesIO(blob)) as zf:
            assert "word/comments.xml" in zf.namelist()
            root = etree.fromstring(zf.read("word/comments.xml"))
            comments = root.findall(f"{{{W_NS}}}comment")
            assert len(comments) == 2

    def test_comment_ids_match_across_xml(self, synthetic_docx: Path) -> None:
        """Every comment id appears as range-start, range-end, and reference."""
        blob = annotate(str(synthetic_docx), _two_finding_report())
        with zipfile.ZipFile(BytesIO(blob)) as zf:
            doc_xml = zf.read("word/document.xml")
            comments_xml = zf.read("word/comments.xml")

        doc_root = etree.fromstring(doc_xml)
        com_root = etree.fromstring(comments_xml)

        comment_ids = _ids_at(com_root, f".//{{{W_NS}}}comment")
        starts = _ids_at(doc_root, f".//{{{W_NS}}}commentRangeStart")
        ends = _ids_at(doc_root, f".//{{{W_NS}}}commentRangeEnd")
        refs = _ids_at(doc_root, f".//{{{W_NS}}}commentReference")

        assert comment_ids == starts == ends == refs
        assert len(comment_ids) == 2

    def test_comment_text_carries_severity_and_rule_id(
        self, synthetic_docx: Path
    ) -> None:
        blob = annotate(str(synthetic_docx), _two_finding_report())
        with zipfile.ZipFile(BytesIO(blob)) as zf:
            text = zf.read("word/comments.xml").decode("utf-8")
        assert "[error]" in text
        assert "[warning]" in text
        assert "layout.margins" in text
        assert "structure.abstract_length" in text

    def test_pass_results_skipped(self, synthetic_docx: Path) -> None:
        report = CheckReport(
            file_name="x.docx",
            timestamp=datetime.now(timezone.utc),
            summary=ReportSummary(
                total_checks=1, passed=1, errors=0, warnings=0, info=0
            ),
            results=[
                CheckResult(
                    rule_id="layout.margins",
                    category="Layout",
                    name="Margin check",
                    status="pass",
                    severity="error",
                    details=[],
                ),
            ],
        )
        blob = annotate(str(synthetic_docx), report)
        with zipfile.ZipFile(BytesIO(blob)) as zf:
            if "word/comments.xml" in zf.namelist():
                root = etree.fromstring(zf.read("word/comments.xml"))
                assert root.findall(f"{{{W_NS}}}comment") == []

    def test_out_of_range_index_falls_back(self, synthetic_docx: Path) -> None:
        """An anchor pointing past the doc must not raise; fall back to first para."""
        report = CheckReport(
            file_name="x.docx",
            timestamp=datetime.now(timezone.utc),
            summary=ReportSummary(
                total_checks=1, passed=0, errors=1, warnings=0, info=0
            ),
            results=[
                CheckResult(
                    rule_id="layout.margins",
                    category="Layout",
                    name="Margin check",
                    status="fail",
                    severity="error",
                    details=[
                        CheckDetail(
                            location="paragraph 99999",
                            locator=Locator(kind="paragraph", paragraph_index=99999),
                            message="x",
                        ),
                    ],
                ),
            ],
        )
        blob = annotate(str(synthetic_docx), report)
        Document(BytesIO(blob))

    def test_reannotating_appends_without_id_collision(
        self, synthetic_docx: Path, tmp_path: Path
    ) -> None:
        """
        Annotating a docx that already has comments must append fresh ids
        instead of overwriting. Guards against the XmlPart/_blob trap where
        python-docx loads /word/comments.xml as ``CommentsPart`` and writes
        from ``_element`` on save.
        """
        report = _two_finding_report()
        first = annotate(str(synthetic_docx), report)
        intermediate = tmp_path / "first.docx"
        intermediate.write_bytes(first)

        second = annotate(str(intermediate), report)

        with zipfile.ZipFile(BytesIO(second)) as zf:
            root = etree.fromstring(zf.read("word/comments.xml"))

        raw_ids = [
            c.get(f"{{{W_NS}}}id") for c in root.findall(f"{{{W_NS}}}comment")
        ]
        ids = sorted(int(x) for x in raw_ids if x is not None)
        assert ids == [0, 1, 2, 3]

    def test_comment_reference_style_added(self, synthetic_docx: Path) -> None:
        """The exporter must add CommentReference style to styles.xml."""
        blob = annotate(str(synthetic_docx), _two_finding_report())
        with zipfile.ZipFile(BytesIO(blob)) as zf:
            styles_xml = zf.read("word/styles.xml")
        root = etree.fromstring(styles_xml)
        for s in root.iterfind(f".//{{{W_NS}}}style"):
            if s.get(f"{{{W_NS}}}styleId") == "CommentReference":
                break
        else:
            pytest.fail("CommentReference style was not added to styles.xml")
